# -------------------------------------------------------------------------------------------------------------------- #
#                                                        IMPORTS                                                       #
# -------------------------------------------------------------------------------------------------------------------- #

#standard library imports
import sys, os, re, shutil, threading
from datetime import datetime

#third party imports
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
#from docx2pdf import convert as convert2pdf
from comtypes.client import CreateObject

#local imports
from tooltip import CreateToolTip

#changing the working directory, so that the program can be run from any path
os.chdir(os.path.dirname(os.path.realpath(__file__)))

# -------------------------------------------------------------------------------------------------------------------- #
#                                                 WINDOW INITIALIZATION                                                #
# -------------------------------------------------------------------------------------------------------------------- #

#window creation
window = tk.Tk()
window.title("Musikredaktion Rotationsmanager")
window.minsize(485, 305)

#tkinter variables
kalenderwoche = tk.StringVar(window, datetime.now().isocalendar()[1])
fortschritt = tk.IntVar(window, 0)
pdf_export = tk.BooleanVar(window, False)

#keyboard bindings
def focus_next_widget(event):
	event.widget.tk_focusNext().focus()
	return("break")
def focus_prev_widget(event):
	event.widget.tk_focusPrev().focus()
	return("break")
window.bind("<Tab>", focus_next_widget)
window.bind("<Shift-Tab>", focus_prev_widget)

#theme: https://github.com/rdbende/Azure-ttk-theme
window.tk.call("source", os.path.join("theme", "azure.tcl"))
#window.tk.call("set_theme", "light")
window.tk.call("set_theme", "dark")

def tkinter_center(win: tk.Tk | tk.Toplevel):
	"""Centers a tkinter window on the screen.
	Copied from https://stackoverflow.com/a/10018670"""
	win.update_idletasks()
	width = win.winfo_width()
	frm_width = win.winfo_rootx() - win.winfo_x()
	win_width = width + 2 * frm_width
	height = win.winfo_height()
	titlebar_height = win.winfo_rooty() - win.winfo_y()
	win_height = height + titlebar_height + frm_width
	x = win.winfo_screenwidth() // 2 - win_width // 2
	y = win.winfo_screenheight() // 2 - win_height // 2
	win.geometry(f"{width}x{height}+{x}+{y}")
	win.deiconify()

# -------------------------------------------------------------------------------------------------------------------- #
#                                                       COMMANDS                                                       #
# -------------------------------------------------------------------------------------------------------------------- #

filetypes_txt:  list[tuple[str, str]] = [("Textdatei", "*.txt"),  ("Alle Dateien", "*.*")]
filetypes_pdf:  list[tuple[str, str]] = [("PDF-Datei", "*.pdf"),  ("Alle Dateien", "*.*")]
filetypes_docx: list[tuple[str, str]] = [("Worddatei", "*.docx"), ("Alle Dateien", "*.*")]

def get_calendar_week_string() -> str:
	"""Returns the calendar week as a formatted string."""
	return f"KW {kalenderwoche.get()}"

def import_file():
	"""Imports the text from a .txt file when the user clicks the import button."""
	try:
		filename = filedialog.askopenfilename(filetypes = filetypes_txt, title = "Textdatei importieren")
		if filename is not None and filename != "" and os.path.isfile(filename):
			text.delete(1.0, tk.END)
			with open(filename, "r", encoding = "utf-8") as file:
				strings = file.readlines()
			text.insert(tk.END, "\n".join(filter(None, (string.strip() for string in strings))))
	except Exception as e:
		messagebox.showerror("Fehler beim Textdatei importieren", e)

def export_file():
	"""Exports the text to a .txt file when the user clicks the export button."""
	try:
		filename = filedialog.asksaveasfilename(confirmoverwrite = True, defaultextension = ".txt", filetypes = filetypes_txt, initialfile = get_calendar_week_string(), title = "Textdatei exportieren")
		if filename is not None and filename != "":
			with open(filename, "w", encoding = "utf-8") as file:
				file.write(text.get("1.0", tk.END))
	except Exception as e:
		messagebox.showerror("Fehler beim Textdatei exportieren", e)
	finally:
		if file and file is not None and not file.closed:
			file.close()

def sort_lines():
	"""Sorts the lines in the text widget alphabetically."""
	lines = text.get("1.0", tk.END).split("\n")
	lines = filter(None, (line.strip() for line in lines))
	#known bug: english and german articles in the title and comments are also ignored, not just in the artist name
	lines = sorted(lines, key = lambda s: re.sub("(\A\s*|(?<=\s))(the|der|die|das)(\s*\Z|\s+(?=\w))", "", s, flags = re.IGNORECASE).lower())
	text.delete(1.0, tk.END)
	text.insert(tk.END, "\n".join(lines))

def get_artist_and_title(string: str) -> str:
	index = string.find("--")
	return string if index == -1 else string[:index-1]

def generate_docx() -> str | None:
	#open template file and check if it contains at least one table
	try:
		filename = filedialog.asksaveasfilename(confirmoverwrite = True, defaultextension = ".docx", filetypes = filetypes_docx, initialfile = get_calendar_week_string(), title = "Worddatei exportieren")
		if filename == "" or not os.path.isdir(os.path.dirname(filename)):
			return None
	except Exception as e:
		messagebox.showerror("Fehler beim Worddatei exportieren", e)
	#load template
	document = Document("template.docx")
	if len(document.tables) == 0:
		messagebox.showerror("Error", "template.docx muss mindestens eine Tabelle enthalten.")
		return None
	#replace the calendar week in the document
	for p in document.paragraphs:
		if "Playlisten-Rotation: KW" in p.text:
			for run in p.runs:
				if "KW" in run.text:
					run.text = run.text.replace("KW", get_calendar_week_string())
					break
			break
	#create generator for the processed strings from the text widget
	line_generator = (string.strip() for string in text.get("1.0", tk.END).split("\n"))
	line_generator = map(get_artist_and_title, filter(None, line_generator))
	table = document.tables[0]
	#fill the table and dynamically add rows
	#ressource: https://python-docx.readthedocs.io/en/latest/api/table.html
	table.cell(0, 0).text = next(line_generator, "")
	for line in line_generator:
		table.add_row().cells[0].text = line
	#save the document
	document.save(filename)
	return filename

#disable PDF export on platforms where it is not implemented
pdf_implemented_platforms = ["win32"]
cb_state = tk.NORMAL
cb_cursor = "hand2"
if sys.platform not in pdf_implemented_platforms:
	pdf_export.set(False)
	cb_state = tk.DISABLED
	cb_cursor = "X_cursor"

def update_progressbar(value: int):
	try:
		fortschritt.set(value)
		window.update()
	except RuntimeError:
		pass

def generate_pdf_of_diff_platforms(docx_filename: str, pdf_filename: str):
	#filename = generate_docx_wrapper()
	#if filename is not None:
	#	convert2pdf(filename, filename.replace(".docx", ".pdf"))
	#https://github.com/python-openxml/python-docx/issues/113#issuecomment-66320799
	#https://stackoverflow.com/questions/6011115/doc-to-pdf-using-python
	#for windows:
	if sys.platform == "win32":
		#launch Microsoft Word
		try:
			word = CreateObject("Word.Application")
			word.Visible = False
		except Exception as e:
			messagebox.showerror("Word nicht installiert", "Microsoft Word ist nicht installiert. " + e)
			return
		update_progressbar(3)
		try:
			#open the docx file
			worddoc = word.Documents.Open(os.path.abspath(docx_filename))
			update_progressbar(5)
			#save document as pdf
			worddoc.SaveAs(os.path.abspath(pdf_filename), FileFormat = 17)
			update_progressbar(7)
			worddoc.Close()
			update_progressbar(8)
		except Exception as e:
			messagebox.showerror("Fehler beim PDF-Datei exportieren", e)
		finally:
			word.Quit()
			update_progressbar(9)
	#for macos:
	elif sys.platform == "darwin":
		pass
	#for linux and other platforms:
	else:
		if shutil.which("libreoffice") is not None:
			pass

def generate_pdf(docx_filename: str):
	"""Generates a .docx and .pdf file from the text input when the user clicks the corresponding button."""
	#get save location for the pdf file
	update_progressbar(1)
	try:
		pdf_filename = filedialog.asksaveasfilename(confirmoverwrite = True, defaultextension = ".pdf", filetypes = filetypes_pdf, initialfile = get_calendar_week_string(), title = "PDF-Datei exportieren")
		if pdf_filename == "" or not os.path.isdir(os.path.dirname(pdf_filename)):
			return None
	except Exception as e:
		messagebox.showerror("Fehler beim PDF-Datei exportieren", e)
		return
	update_progressbar(2)

	#launch thread so that the GUI does not freeze
	thread = threading.Thread(target = generate_pdf_of_diff_platforms, args = (docx_filename, pdf_filename), daemon = False)
	thread.start()

def generate():
	update_progressbar(0)
	filename = generate_docx()
	if filename is None:
		return
	if pdf_export.get():
		generate_pdf(filename)
	else:
		update_progressbar(9)

# -------------------------------------------------------------------------------------------------------------------- #
#                                                        LAYOUT                                                        #
# -------------------------------------------------------------------------------------------------------------------- #

padding = 12

frame_left = tk.Frame(window)

#text input with scrollbar
frame_text = tk.Frame(frame_left)
scrollbar = ttk.Scrollbar(frame_text)
text = tk.Text(frame_text, width = 40, height = 15, undo = True, yscrollcommand = scrollbar.set)
scrollbar.config(command = text.yview)
scrollbar.pack(side = tk.RIGHT, fill = tk.Y)
text.pack(fill = tk.BOTH, expand = True)
frame_text.pack(fill = tk.BOTH, expand = True)

ttk.Label(frame_left, text = "Version 2 © 2024, Jona Heinke unter MIT Lizenz").pack(side = tk.BOTTOM, padx = padding, pady = padding)

frame_left.pack(side = tk.LEFT, fill = tk.BOTH, expand = True)

frame_controls = tk.Frame(window)

frame_controls_top = tk.Frame(frame_controls)
button = ttk.Button(frame_controls_top, text = "Importieren ⭳", cursor = "bottom_side", command = import_file)
CreateToolTip(button, "Importiert den Text aus einer Textdatei.")
button.pack(fill = tk.X)
button = ttk.Button(frame_controls_top, text = "Exportieren ↥", cursor = "top_side", command = export_file)
CreateToolTip(button, "Exportiert den Text in eine Textdatei.")
button.pack(fill = tk.X, pady = padding)
button = ttk.Button(frame_controls_top, text = "Sortieren 🗘", cursor = "exchange", command = sort_lines)
CreateToolTip(button, "Sortiert alle Lieder alphabetisch. \"the\", \"der\", \"die\" und \"das\" werden dabei ignoriert.")
button.pack(fill = tk.X)
button = ttk.Button(frame_controls_top, text = "DOCX generieren →", cursor = "right_side", command = generate, style = "Accent.TButton")
CreateToolTip(button, "Generiert die Worddatei aus dem eingegebenen Text.")
button.pack(fill = tk.X, pady = (padding, 0))
ttk.Progressbar(frame_controls_top, value = 0, maximum = 9, variable = fortschritt, mode = "determinate").pack(fill = tk.X, pady = (padding, 0))
frame_controls_top.pack(side = tk.TOP, fill = tk.X)

frame_controls_bottom = tk.Frame(frame_controls)
cb = ttk.Checkbutton(frame_controls_bottom, text = "PDF exportieren", cursor = cb_cursor, variable = pdf_export, style = "Switch.TCheckbutton", state = cb_state)
CreateToolTip(cb, "Erstellt zusätzlich eine PDF Datei, wenn die Worddatei generiert wird.\nHINWEIS: dauert bis zu zwanzig Sekunden.")
cb.pack(pady = (0, padding))
ttk.Label(frame_controls_bottom, text = "Kalenderwoche:").pack(side = tk.LEFT, padx = (0, 5))
ttk.Spinbox(frame_controls_bottom, from_ = 1, to = 52, textvariable = kalenderwoche, width = 5).pack()
frame_controls_bottom.pack(side = tk.BOTTOM, fill = tk.X)

frame_controls.pack(side = tk.RIGHT, fill = tk.BOTH, padx = padding, pady = padding)

# -------------------------------------------------------------------------------------------------------------------- #
#                                                       MAIN LOOP                                                      #
# -------------------------------------------------------------------------------------------------------------------- #

tkinter_center(window)
window.mainloop()